import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = None # default black
    
    # Set paragraph formatting
    p_format = style.paragraph_format
    p_format.line_spacing = 1.15
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading_front(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_chapter_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Title case logic (capitalize each word)
    capitalized = " ".join([word.capitalize() for word in text.split(" ")])
    run = p.add_run(capitalized)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_code(doc, code_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_str)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

def main():
    doc = Document()
    set_style(doc)

    # FRONT MATTER
    front_matter = [
        "Cover Page", "Declaration", "Dedication", "Acknowledgment", 
        "Table of Contents", "Abstract", "List of Figures", 
        "List of Tables", "Definition of Key Terms", "Abbreviations and Acronyms"
    ]
    for fm in front_matter:
        add_heading_front(doc, fm)
        doc.add_paragraph(f"[Content for {fm}]")
        doc.add_page_break()

    # CHAPTER ONE
    add_chapter_heading(doc, "CHAPTER ONE: SYSTEM OVERVIEW")
    
    add_subheading(doc, "1.1 Research Problem")
    doc.add_paragraph("Food waste is a growing global concern, with household spoilage being a major contributor. Often, individuals purchase groceries and forget about them until they pass their expiration dates. This leads to massive economic losses and an increased environmental footprint, as perfectly consumable food is discarded.")
    
    add_subheading(doc, "1.2 System Objectives")
    add_subheading(doc, "1.2.1 General Objective")
    doc.add_paragraph("The general objective of this project is to develop a web-based Food Expiry Tracker application that allows users to monitor the shelf life of their groceries, receive alerts, and ultimately minimize household food waste.")
    add_subheading(doc, "1.2.2 Specific Objectives")
    doc.add_paragraph("1. To provide an intuitive user interface for adding and categorizing food items.\n2. To track the expiry dates of items and flag those that are expiring soon.\n3. To offer recipe suggestions based on currently available and expiring ingredients.\n4. To provide an administrative dashboard for monitoring system usage and managing users.")

    add_subheading(doc, "1.3 System Scope")
    doc.add_paragraph("The system focuses on individual households and personal pantry management. It encompasses a frontend web application for user interactions and a backend SQLite database for secure data persistence.")

    add_subheading(doc, "1.4 System Justification")
    doc.add_paragraph("By empowering users with tools to actively track food expiration dates, this system will directly contribute to lowering food waste. This results in significant cost savings for the user and promotes more sustainable consumption patterns.")
    doc.add_page_break()

    # CHAPTER TWO
    add_chapter_heading(doc, "CHAPTER TWO: FRONT END (USER INTERFACE) IMPLEMENTATION")
    
    add_subheading(doc, "2.1 Home Page")
    doc.add_paragraph("[Insert Screenshot of the Home Page here]")
    add_caption(doc, "Fig 2.1-1 Home Page Interface")
    doc.add_paragraph("The Home Page UI generator code includes the following HTML structure:")
    add_code(doc, '''<!-- Landing Page -->
<div class="landing-page-container" id="landing-page-container">
    <header class="landing-header">
        <div class="logo">
            <i class="fa-solid fa-leaf"></i> Freshr
        </div>
        <nav class="nav-links">
            <button id="btn-landing-login">Log In</button>
            <button class="primary" id="btn-landing-signup">Sign Up</button>
        </nav>
    </header>
</div>''')

    add_subheading(doc, "2.2 Products Page")
    doc.add_paragraph("[Insert Screenshot of the Products Page here]")
    add_caption(doc, "Fig 2.2-1 Products Page Interface")
    doc.add_paragraph("The Products Page UI generator code includes the following HTML structure:")
    add_code(doc, '''<!-- App Main Content -->
<div class="app-container" id="app-container" style="display:none;">
    <main class="main-content">
        <header class="top-nav">
            <div class="search-bar">
                <i class="fa-solid fa-search"></i>
                <input type="text" id="search-input" placeholder="Search your food...">
            </div>
        </header>
        <div class="inventory-grid" id="inventory-grid-container">
            <!-- Dynamic Food Cards Go Here -->
        </div>
    </main>
</div>''')
    doc.add_page_break()

    # CHAPTER THREE
    add_chapter_heading(doc, "CHAPTER THREE: BACK END (LOGIC) IMPLEMENTATION")
    
    add_subheading(doc, "3.1 System Processes")
    add_subheading(doc, "3.1.1 Login Process")
    doc.add_paragraph("Pseudocode for Login:")
    add_code(doc, '''BEGIN
    RECEIVE username and password from user
    IF username OR password is empty THEN
        RETURN Error "Credentials required"
    END IF
    FETCH user_record FROM database WHERE username = username
    IF user_record exists THEN
        COMPARE password WITH user_record.password_hash
        IF match THEN
            GENERATE auth_token
            RETURN auth_token and Success
        ELSE
            RETURN Error "Authentication failed"
        END IF
    ELSE
        RETURN Error "Authentication failed"
    END IF
END''')
    doc.add_paragraph("[Insert Flowchart of Login Process here]")
    add_caption(doc, "Fig 3.1.1-1 Login Process Flowchart")
    doc.add_paragraph("[Insert Screenshot of Login Code here]")
    add_caption(doc, "Fig 3.1.1-2 Login Code snippet")
    doc.add_paragraph("Backend login implementation code:")
    add_code(doc, '''app.post('/api/auth/signin', async (req, res) => {
    const { username, password } = req.body;
    if (!username || !password) return res.status(400).json({ error: 'Required' });
    try {
        const user = await dbGet('SELECT * FROM users WHERE username = ?', [username]);
        if (!user) return res.status(401).json({ error: 'Failed' });
        const match = await bcrypt.compare(password, user.password);
        if (!match) return res.status(401).json({ error: 'Failed' });
        const token = jwt.sign({ username, role: user.role }, JWT_SECRET);
        res.json({ token, user: { username, role: user.role } });
    } catch (err) {
        res.status(500).json({ error: 'Server error' });
    }
});''')

    add_subheading(doc, "3.1.2 Add Product Process")
    doc.add_paragraph("Pseudocode for Adding a Product:")
    add_code(doc, '''BEGIN
    RECEIVE product details (name, category, qty, expiryDate)
    VERIFY user is authenticated
    VALIDATE product details
    IF validation fails THEN
        RETURN Error
    END IF
    INSERT INTO food_items database table
    RETURN Success message
END''')
    doc.add_paragraph("[Insert Flowchart of Add Product Process here]")
    add_caption(doc, "Fig 3.1.2-1 Add Product Process Flowchart")
    doc.add_paragraph("[Insert Screenshot of Add Product Code here]")
    add_caption(doc, "Fig 3.1.2-2 Add Product Code snippet")

    add_subheading(doc, "3.2 Database Implementation")
    add_subheading(doc, "3.2.1 Normalization (1nf, 2nf, 3nf)")
    doc.add_paragraph("1NF: The raw data is organized into tables (users, food_items) where each column contains atomic values, and each record is unique via a primary key.")
    doc.add_paragraph("2NF: The tables meet 1NF and all non-key attributes are fully functionally dependent on the primary key. For example, food item details depend entirely on the food item ID.")
    doc.add_paragraph("3NF: The tables meet 2NF and contain no transitive dependencies. User roles and details are stored in the users table independently of the food items they add.")
    
    add_subheading(doc, "3.2.2 Database Technology")
    doc.add_paragraph("SQLite was used for the database management system. It is a lightweight, zero-configuration database that is perfectly suited for applications needing a local, serverless storage solution while providing full relational database features.")

    add_subheading(doc, "3.2.3 Database Tables")
    doc.add_paragraph("[Insert Screenshot of users table here]")
    add_caption(doc, "Table 3.2.3-1 Users Table")
    doc.add_paragraph("[Insert Screenshot of food_items table here]")
    add_caption(doc, "Table 3.2.3-2 Food Items Table")
    doc.add_paragraph("[Insert Screenshot of history_log table here]")
    add_caption(doc, "Table 3.2.3-3 History Log Table")
    
    add_subheading(doc, "3.2.4 System-to-database Integration")
    doc.add_paragraph("The application communicates with the database using parameterized queries to ensure security against SQL injection.")
    add_code(doc, '''// Sign up Database Integration
const hash = await bcrypt.hash(password, 10);
await dbRun('INSERT INTO users (username, password, role, provider) VALUES (?, ?, ?, ?)', [username, hash, role, 'local']);

// Add Food Database Integration
await dbRun(
    `INSERT INTO food_items (id, username, name, category, storage, qty, unit, dateAdded, dateExpiry, imageData)
     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)`,
    [id, req.user.username, name, category, storage, qty, unit, dateAdded, dateExpiry, imageData || null]
);''')
    doc.add_page_break()

    # CHAPTER FOUR
    add_chapter_heading(doc, "CHAPTER FOUR: CONCLUSION AND RECOMMENDATIONS")
    
    add_subheading(doc, "4.1 Conclusion")
    doc.add_paragraph("In conclusion, this project successfully implements a comprehensive web-based platform for food expiry tracking. Users can seamlessly create accounts, log in, and manage their pantry by adding, viewing, and categorizing food products. Furthermore, the system includes an administrative dashboard that enables the shop owner or admin to monitor users and activity logs. The user interface provides clear error and success messages, along with interactive prompts, ensuring a smooth and user-friendly experience throughout.")

    add_subheading(doc, "4.2 Recommendation")
    doc.add_paragraph("Based on the experience of developing this project, a future developer would be advised to integrate automated data entry methods. While the current system efficiently handles manual input, integrating barcode scanning or Optical Character Recognition (OCR) for supermarket receipts could drastically reduce user friction. Technologies like the QuaggaJS library for barcodes or Tesseract.js for OCR could be easily incorporated into the frontend. Additionally, migrating from SQLite to a scalable cloud database (such as PostgreSQL on AWS) would be recommended if the application is to be deployed for a massive, multi-tenant user base.")

    output_path = os.path.join(os.getcwd(), 'Project_Documentation.docx')
    doc.save(output_path)
    print(f"Documentation saved to {output_path}")

if __name__ == "__main__":
    main()
